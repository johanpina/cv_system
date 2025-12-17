import os
import sys
import io
import re
import time
import logging
import zipfile
from typing import Optional, Tuple
from tqdm import tqdm
from docx import Document 

# Configuración de Paths
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from sqlmodel import Session, select
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import google.generativeai as genai
from google.api_core import exceptions as google_exceptions
from googleapiclient.errors import HttpError

from app.core.database import engine
from app.models.models import Url_HojaDeVida
from app.core.config import settings

# --- CONFIGURACIÓN DE LOGS ---
logging.basicConfig()
logging.getLogger('sqlalchemy.engine').setLevel(logging.WARNING)
logging.getLogger('googleapiclient.discovery_cache').setLevel(logging.ERROR)

# --- CONFIGURACIÓN ---
SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
SERVICE_ACCOUNT_FILE = 'scripts/credentials.json'

genai.configure(api_key=settings.GOOGLE_API_KEY)

def get_drive_service():
    try:
        creds = service_account.Credentials.from_service_account_file(
            SERVICE_ACCOUNT_FILE, scopes=SCOPES)
        return build('drive', 'v3', credentials=creds)
    except Exception as e:
        print(f"❌ Error crítico autenticando Drive: {e}")
        return None

def extract_id_from_url(url: str) -> Optional[str]:
    if not url: return None
    patterns = [r'id=([a-zA-Z0-9_-]+)', r'/d/([a-zA-Z0-9_-]+)']
    for pattern in patterns:
        match = re.search(pattern, url)
        if match: return match.group(1)
    return None

def convert_docx_bytes_to_text(docx_bytes: bytes) -> Optional[bytes]:
    """
    Intenta convertir DOCX (o DOCM) a texto plano.
    """
    try:
        doc = Document(io.BytesIO(docx_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_text))

        return "\n".join(full_text).encode('utf-8')

    except Exception:
        # Si falla (ej: es un .doc renombrado), retornamos None
        return None

def smart_download_file(service, file_id: str) -> Optional[Tuple[bytes, str]]:
    """
    Descarga SOLO formatos soportados. Filtra los .doc viejos para evitar errores 400.
    """
    try:
        # 1. Obtener Metadatos
        file_metadata = service.files().get(fileId=file_id, fields="name, mimeType, size").execute()
        original_mime = file_metadata.get('mimeType')
        
        request = None
        target_mime = None
        needs_conversion_to_text = False

        # --- LISTA BLANCA ESTRICTA ---
        
        # CASO 1: Google Doc -> Exportar a PDF
        if original_mime == 'application/vnd.google-apps.document':
            request = service.files().export(fileId=file_id, mimeType='application/pdf')
            target_mime = 'application/pdf'
            
        # CASO 2: PDF Nativo -> Descargar directo
        elif original_mime == 'application/pdf':
            request = service.files().get_media(fileId=file_id)
            target_mime = 'application/pdf'

        # CASO 3: Word Moderno (.docx) O con Macros (.docm) -> Descargar y Convertir a Texto
        elif original_mime in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document', # .docx
            'application/vnd.ms-word.document.macroenabled.12' # .docm
        ]:
            request = service.files().get_media(fileId=file_id)
            target_mime = 'text/plain' # Gemini recibirá texto plano
            needs_conversion_to_text = True

        # CASO 4: Todo lo demás (incluido application/msword que son los .doc viejos) -> IGNORAR
        else:
            # print(f"⚠️ Formato no soportado ignorado: {original_mime}")
            return None

        # 2. Ejecutar Descarga
        file_stream = io.BytesIO()
        downloader = MediaIoBaseDownload(file_stream, request)
        done = False
        while done is False:
            status, done = downloader.next_chunk()
        
        data = file_stream.getvalue()

        # 3. Conversión (Si aplica)
        if needs_conversion_to_text:
            text_data = convert_docx_bytes_to_text(data)
            if text_data is None:
                # Falló la conversión (probablemente archivo corrupto)
                return None 
            data = text_data

        # 4. Validación de contenido vacío
        if len(data) < 50: 
            return None
            
        return data, target_mime

    except HttpError:
        return None
    except Exception:
        return None

def analyze_cv_with_gemini(file_content: bytes, mime_type: str) -> Optional[str]:
    model = genai.GenerativeModel("models/gemini-2.0-flash") 
    
    prompt = """
    Analiza la información del CV proporcionado.
    SI EL DOCUMENTO ESTÁ VACÍO, ILEGIBLE O NO ES UN CV, RESPONDE: "DATOS_NO_DISPONIBLES".
    
    De lo contrario, extrae el resumen siguiendo ESTRICTAMENTE este formato plano:

    PERFIL_PROFESIONAL: [Resumen]
    TITULOS_ACADEMICOS: [Lista]
    HABILIDADES_TECNICAS: [Lista]
    EXPERIENCIA_DOCENTE: [Resumen]
    EXPERIENCIA_INDUSTRIA: [Resumen]
    IDIOMAS: [Idiomas]
    """
    
    try:
        response = model.generate_content([
            {'mime_type': mime_type, 'data': file_content},
            prompt
        ])
        
        text = response.text
        if "DATOS_NO_DISPONIBLES" in text:
            return None
            
        return text

    except google_exceptions.ResourceExhausted:
        print("\n⚠️  Cuota excedida. Pausando 20s...")
        time.sleep(20)
        return None
    except Exception as e:
        # print(f"\n❌ Error Gemini: {e}") 
        return None

def main():
    print("🚀 Iniciando Motor (Filtro Inteligente: Solo PDF/DOCX/DOCM)")
    
    drive_service = get_drive_service()
    if not drive_service: return

    with Session(engine) as session:
        statement = select(Url_HojaDeVida).where(Url_HojaDeVida.resumen_estructurado == None)
        cvs = session.exec(statement).all()
        
        total_cvs = len(cvs)
        print(f"📊 Pendientes: {total_cvs}")
        
        pbar = tqdm(cvs, desc="Procesando", unit="cv")
        processed_count = 0
        
        for cv in pbar:
            try:
                file_id = extract_id_from_url(cv.url_hoja_de_vida)
                if not file_id: continue

                result = smart_download_file(drive_service, file_id)
                
                if result:
                    file_data, mime_type = result 
                    
                    # Log visual limpio
                    label = "TXT" if "text" in mime_type else "PDF"
                    pbar.set_description(f"ID {cv.id_aspirante} ({label} -> AI)")
                    
                    resumen = analyze_cv_with_gemini(file_data, mime_type)
                    
                    if resumen:
                        cv.resumen_estructurado = resumen
                        session.add(cv)
                        session.commit()
                        processed_count += 1
                        pbar.set_description(f"ID {cv.id_aspirante} ✅")
                    else:
                        session.rollback()
                else:
                    # Si entra aquí es porque el archivo era .doc viejo o corrupto.
                    # No hacemos nada, solo pasamos al siguiente.
                    pass
                
                time.sleep(1)

            except Exception:
                session.rollback()
                continue 

    print(f"\n🏁 Finalizado. Éxito: {processed_count}/{total_cvs}")

if __name__ == "__main__":
    main()